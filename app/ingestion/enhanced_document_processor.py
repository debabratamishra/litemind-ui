import fitz
import pandas as pd
import numpy as np
import cv2
import base64
import logging
from PIL import Image
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import re
from collections import defaultdict
import easyocr
from docx import Document
from docx.shared import Inches
import zipfile
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
import pdfplumber
import camelot
import warnings

logging.getLogger('pypdf._reader').setLevel(logging.ERROR)
logging.getLogger('pypdf').setLevel(logging.ERROR)

warnings.filterwarnings('ignore', module='pypdf')

logger = logging.getLogger(__name__)

class EnhancedDocumentProcessor:
    """Enhanced processor for PDFs and mixed-content documents."""
    
    def __init__(self):
        try:
            # Initialize OCR reader for fallback text extraction
            self.ocr_reader = easyocr.Reader(['en'], gpu=False)
            self.ocr_available = True
        except Exception as e:
            logger.warning(f"OCR initialization failed: {e}")
            self.ocr_reader = None
            self.ocr_available = False
        
        # Document analysis settings
        self.min_image_size = (100, 100)  # Minimum image dimensions
        self.max_chunk_size = 1000
        self.overlap_size = 100
        
    def extract_pdf_enhanced(self, path: Path) -> List[dict]:
        """Enhanced PDF extraction with layout analysis and intelligent chunking."""
        try:
            logger.info(f"Starting enhanced PDF processing for {path.name}")
            
            # Try multiple extraction methods for best results
            content_blocks = []
            
            # Method 1: PyMuPDF for general extraction
            pymupdf_content = self._extract_with_pymupdf(path)
            content_blocks.extend(pymupdf_content)
            
            # Method 2: pdfplumber for table detection
            table_content = self._extract_tables_with_pdfplumber(path)
            content_blocks.extend(table_content)
            
            # Method 3: Camelot for complex tables (if available)
            try:
                camelot_tables = self._extract_tables_with_camelot(path)
                content_blocks.extend(camelot_tables)
            except Exception as e:
                logger.debug(f"Camelot table extraction failed: {e}")
            
            # Method 4: Document structure analysis
            structure_analysis = self._analyze_document_structure(path)
            if structure_analysis:
                content_blocks.append(structure_analysis)
            
            logger.info(f"Extracted {len(content_blocks)} content blocks from PDF")
            return content_blocks
            
        except Exception as e:
            logger.error(f"Enhanced PDF extraction failed for {path}: {e}")
            return self._create_fallback_pdf_content(path, str(e))
    
    def _analyze_page_layout(self, page, page_num: int, filename: str) -> Optional[dict]:
        """Analyze page layout and structure."""
        try:
            # Get page dimensions
            rect = page.rect
            page_width = rect.width
            page_height = rect.height
            
            # Get text blocks for layout analysis
            text_dict = page.get_text("dict")
            
            # Analyze layout characteristics
            layout_info = {
                "page_dimensions": f"{page_width:.1f}x{page_height:.1f}",
                "text_blocks": len(text_dict.get("blocks", [])),
                "images": len(page.get_images()),
                "estimated_reading_time": 0,
                "layout_type": "unknown"
            }
            
            # Calculate text density and reading time
            full_text = page.get_text()
            word_count = len(full_text.split())
            layout_info["word_count"] = word_count
            layout_info["estimated_reading_time"] = max(1, word_count // 200)  # ~200 words per minute
            
            # Determine layout type based on content
            if layout_info["images"] > 3 and word_count < 500:
                layout_info["layout_type"] = "image_heavy"
            elif layout_info["text_blocks"] > 10 and word_count > 1000:
                layout_info["layout_type"] = "text_dense"
            elif layout_info["images"] > 0 and word_count > 200:
                layout_info["layout_type"] = "mixed_content"
            elif word_count < 100:
                layout_info["layout_type"] = "minimal_content"
            else:
                layout_info["layout_type"] = "standard_document"
            
            # Only create layout analysis if there's substantial content
            if word_count > 50 or layout_info["images"] > 0:
                content = f"""PAGE LAYOUT ANALYSIS - Page {page_num} of {filename}:

                Layout Characteristics:
                • Page Size: {layout_info['page_dimensions']}
                • Text Blocks: {layout_info['text_blocks']}
                • Images: {layout_info['images']}
                • Word Count: {layout_info['word_count']}
                • Layout Type: {layout_info['layout_type']}
                • Estimated Reading Time: {layout_info['estimated_reading_time']} minute(s)

                This page appears to be a {layout_info['layout_type']} page with {"substantial" if word_count > 200 else "minimal"} text content.
                """
                
                return {
                    "content": content,
                    "metadata": {
                        "filename": filename,
                        "page_number": page_num,
                        "content_type": "page_layout_analysis",
                        **layout_info
                    }
                }
            
            return None
            
        except Exception as e:
            logger.debug(f"Page layout analysis failed for page {page_num}: {e}")
            return None

    
    def _extract_with_pymupdf(self, path: Path) -> List[dict]:
        """Enhanced PyMuPDF extraction with layout analysis."""
        content_blocks = []
        
        try:
            doc = fitz.open(str(path))
            document_metadata = doc.metadata
            
            # Create document overview
            overview = self._create_document_overview(doc, path.name)
            content_blocks.append(overview)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Enhanced text extraction with layout preservation
                text_content = self._extract_page_text_enhanced(page, page_num + 1, path.name)
                if text_content:
                    content_blocks.extend(text_content)
                
                # Enhanced image extraction with context
                image_content = self._extract_page_images_enhanced(page, page_num + 1, path.name)
                content_blocks.extend(image_content)
                
                # Extract page-level metadata and structure (with error handling)
                try:
                    page_analysis = self._analyze_page_layout(page, page_num + 1, path.name)
                    if page_analysis:
                        content_blocks.append(page_analysis)
                except AttributeError:
                    logger.debug(f"Page layout analysis not available for page {page_num + 1}")
                except Exception as e:
                    logger.debug(f"Page layout analysis failed for page {page_num + 1}: {e}")

                if page_analysis:
                    content_blocks.append(page_analysis)
            
            doc.close()
            
        except Exception as e:
            logger.error(f"PyMuPDF extraction failed: {e}")
        
        return content_blocks
    
    def _extract_page_text_enhanced(self, page, page_num: int, filename: str) -> List[dict]:
        """Enhanced text extraction with layout and structure awareness."""
        text_blocks = []
        
        try:
            # Get text with detailed layout information
            text_dict = page.get_text("dict")
            
            # Analyze text layout and structure
            layout_analysis = self._analyze_text_layout(text_dict)
            
            # Group text blocks by layout characteristics
            text_groups = self._group_text_blocks(text_dict)
            
            for group_type, blocks in text_groups.items():
                if not blocks:
                    continue
                
                combined_text = self._combine_text_blocks(blocks, group_type)
                
                if combined_text and len(combined_text.strip()) > 20:
                    text_blocks.append({
                        "content": combined_text,
                        "metadata": {
                            "filename": filename,
                            "page_number": page_num,
                            "content_type": f"pdf_text_{group_type}",
                            "layout_type": group_type,
                            "block_count": len(blocks),
                            "confidence": self._calculate_text_confidence(blocks)
                        }
                    })
            
            # Extract structured content (headers, paragraphs, lists)
            structured_content = self._extract_structured_text(text_dict, page_num, filename)
            text_blocks.extend(structured_content)
            
        except Exception as e:
            logger.warning(f"Enhanced text extraction failed for page {page_num}: {e}")
        
        return text_blocks
    
    def _analyze_text_layout(self, text_dict: dict) -> dict:
        """Analyze text layout characteristics."""
        analysis = {
            "font_sizes": defaultdict(int),
            "font_families": defaultdict(int),
            "text_blocks": 0,
            "average_font_size": 0,
            "has_columns": False,
            "column_count": 1
        }
        
        all_font_sizes = []
        
        for block in text_dict.get("blocks", []):
            if "lines" not in block:
                continue
                
            analysis["text_blocks"] += 1
            
            for line in block["lines"]:
                for span in line["spans"]:
                    font_size = span.get("size", 0)
                    font_family = span.get("font", "unknown")
                    
                    analysis["font_sizes"][font_size] += 1
                    analysis["font_families"][font_family] += 1
                    all_font_sizes.append(font_size)
        
        if all_font_sizes:
            analysis["average_font_size"] = np.mean(all_font_sizes)
        
        # Detect columns based on text block positions
        analysis["has_columns"], analysis["column_count"] = self._detect_columns(text_dict)
        
        return analysis
    
    def _group_text_blocks(self, text_dict: dict) -> dict:
        """Group text blocks by their characteristics."""
        groups = {
            "header": [],
            "body": [],
            "footer": [],
            "sidebar": [],
            "caption": []
        }
        
        page_height = text_dict.get("height", 800)
        page_width = text_dict.get("width", 600)
        
        # Calculate font size statistics for classification
        font_sizes = []
        for block in text_dict.get("blocks", []):
            if "lines" not in block:
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    font_sizes.append(span.get("size", 0))
        
        if not font_sizes:
            return groups
        
        avg_font_size = np.mean(font_sizes)
        large_font_threshold = avg_font_size * 1.2
        
        for block in text_dict.get("blocks", []):
            if "lines" not in block:
                continue
            
            bbox = block.get("bbox", [0, 0, 0, 0])
            x0, y0, x1, y1 = bbox
            
            # Calculate block characteristics
            block_height = y1 - y0
            block_width = x1 - x0
            center_y = (y0 + y1) / 2
            
            # Get dominant font size in block
            block_font_sizes = []
            for line in block["lines"]:
                for span in line["spans"]:
                    block_font_sizes.append(span.get("size", 0))
            
            avg_block_font_size = np.mean(block_font_sizes) if block_font_sizes else avg_font_size
            
            # Classify block based on position and characteristics
            if center_y < page_height * 0.15:  # Top 15% of page
                if avg_block_font_size > large_font_threshold:
                    groups["header"].append(block)
                else:
                    groups["body"].append(block)
            elif center_y > page_height * 0.85:  # Bottom 15% of page
                groups["footer"].append(block)
            elif x0 > page_width * 0.75:  # Right 25% of page
                groups["sidebar"].append(block)
            elif block_height < page_height * 0.1 and avg_block_font_size < avg_font_size * 0.9:
                groups["caption"].append(block)
            else:
                groups["body"].append(block)
        
        return groups
    
    def _combine_text_blocks(self, blocks: List[dict], group_type: str) -> str:
        """Combine text blocks intelligently based on group type."""
        text_parts = []
        
        for block in blocks:
            block_lines = []
            for line in block.get("lines", []):
                line_spans = []
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if text:
                        line_spans.append(text)
                
                if line_spans:
                    # Join spans with single space, clean up multiple spaces
                    line_text = " ".join(line_spans)
                    line_text = re.sub(r'\s+', ' ', line_text).strip()
                    block_lines.append(line_text)
            
            if block_lines:
                # Join lines appropriately based on content type
                if group_type == "header":
                    # Headers: preserve line breaks but clean spacing
                    block_text = "\n".join(block_lines)
                else:
                    # Body text: join with spaces, preserve paragraph breaks
                    block_text = " ".join(block_lines)
                
                text_parts.append(block_text)
        
        if not text_parts:
            return ""
        
        # Final combination based on group type
        if group_type == "header":
            result = "\n".join(text_parts)
        elif group_type == "body":
            # Join body paragraphs with double newline
            result = "\n\n".join(text_parts)
        else:
            result = "\n".join(text_parts)
        
        # Clean up excessive whitespace
        result = re.sub(r'\n{3,}', '\n\n', result)  # Max 2 consecutive newlines
        result = re.sub(r'[ \t]+', ' ', result)     # Normalize spaces/tabs
        return result.strip()
    
    def _extract_structured_text(self, text_dict: dict, page_num: int, filename: str) -> List[dict]:
        """Extract structured content like headers, lists, etc."""
        structured_content = []
        
        # Detect headers based on font size and formatting
        headers = self._detect_headers(text_dict)
        if headers:
            structured_content.append({
                "content": f"PAGE {page_num} HEADERS:\n" + "\n".join(headers),
                "metadata": {
                    "filename": filename,
                    "page_number": page_num,
                    "content_type": "pdf_headers",
                    "header_count": len(headers)
                }
            })
        
        # Detect lists
        lists = self._detect_lists(text_dict)
        if lists:
            structured_content.append({
                "content": f"PAGE {page_num} LISTS:\n" + "\n".join(lists),
                "metadata": {
                    "filename": filename,
                    "page_number": page_num,
                    "content_type": "pdf_lists",
                    "list_count": len(lists)
                }
            })
        
        return structured_content
    
    def _detect_headers(self, text_dict: dict) -> List[str]:
        """Detect header text based on font characteristics."""
        headers = []
        
        # Calculate average font size
        font_sizes = []
        for block in text_dict.get("blocks", []):
            if "lines" not in block:
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    font_sizes.append(span.get("size", 0))
        
        if not font_sizes:
            return headers
        
        avg_font_size = np.mean(font_sizes)
        header_threshold = avg_font_size * 1.15
        
        for block in text_dict.get("blocks", []):
            if "lines" not in block:
                continue
            
            for line in block["lines"]:
                line_text = ""
                is_header = False
                
                for span in line["spans"]:
                    font_size = span.get("size", 0)
                    font_flags = span.get("flags", 0)
                    text = span.get("text", "").strip()
                    
                    # Check if this looks like a header
                    if (font_size > header_threshold or 
                        font_flags & 2**4 or  # Bold
                        (len(text) < 100 and font_size > avg_font_size)):
                        is_header = True
                    
                    line_text += text + " "
                
                if is_header and line_text.strip() and len(line_text.strip()) > 3:
                    headers.append(line_text.strip())
        
        return headers
    
    def _detect_lists(self, text_dict: dict) -> List[str]:
        """Detect list items based on formatting patterns."""
        lists = []
        
        list_patterns = [
            r'^\s*[\•\-\*\+]\s+',  # Bullet points
            r'^\s*\d+[\.\)]\s+',   # Numbered lists
            r'^\s*[a-zA-Z][\.\)]\s+',  # Lettered lists
            r'^\s*[ivxlcdm]+[\.\)]\s+',  # Roman numerals
        ]
        
        combined_pattern = '|'.join(list_patterns)
        
        for block in text_dict.get("blocks", []):
            if "lines" not in block:
                continue
            
            for line in block["lines"]:
                line_text = ""
                for span in line["spans"]:
                    line_text += span.get("text", "") + " "
                
                line_text = line_text.strip()
                if line_text and re.match(combined_pattern, line_text, re.IGNORECASE):
                    lists.append(line_text)
        
        return lists
    
    def _extract_page_images_enhanced(self, page, page_num: int, filename: str) -> List[dict]:
        """Enhanced image extraction with context and quality analysis."""
        image_content = []
        
        try:
            images = page.get_images(full=True)
            
            for img_index, img in enumerate(images):
                try:
                    # Extract image data
                    xref = img[0]
                    pix = fitz.Pixmap(page.parent, xref)
                    
                    # Quality filtering
                    if pix.width < self.min_image_size[0] or pix.height < self.min_image_size[1]:
                        pix = None
                        continue
                    
                    # Convert to bytes
                    if pix.n > 3:  # Convert CMYK to RGB
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    
                    img_bytes = pix.tobytes("png")
                    pix = None  # Free memory
                    
                    # Analyze image context
                    image_analysis = self._analyze_image_context(page, img, page_num, img_index)
                    
                    # Create enhanced image content
                    image_content.append({
                        "content": f"""ENHANCED PDF IMAGE ANALYSIS:
                        Document: {filename}
                        Page: {page_num}
                        Image: {img_index + 1}
                        Dimensions: {image_analysis.get('width', 'unknown')}x{image_analysis.get('height', 'unknown')}
                        Context: {image_analysis.get('context', 'No surrounding text detected')}
                        Position: {image_analysis.get('position', 'unknown')}
                        Likely Type: {image_analysis.get('image_type', 'unknown')}

                        This image appears in the context of: {image_analysis.get('surrounding_text', 'No text context available')}
                        """,
                        "metadata": {
                            "filename": filename,
                            "page_number": page_num,
                            "image_index": img_index,
                            "content_type": "pdf_image_enhanced",
                            "image_bytes": base64.b64encode(img_bytes).decode('utf-8'),
                            **image_analysis
                        }
                    })
                    
                except Exception as e:
                    logger.warning(f"Failed to process image {img_index} on page {page_num}: {e}")
                    continue
        
        except Exception as e:
            logger.warning(f"Image extraction failed for page {page_num}: {e}")
        
        return image_content
    
    def _analyze_image_context(self, page, img_info, page_num: int, img_index: int) -> dict:
        """Analyze the context around an image."""
        analysis = {
            "width": img_info[2] if len(img_info) > 2 else 0,
            "height": img_info[3] if len(img_info) > 3 else 0,
            "context": "unknown",
            "position": "unknown",
            "image_type": "unknown",
            "surrounding_text": ""
        }
        
        try:
            # Get image position from the page
            image_rects = page.get_image_rects(img_info[0])
            
            if image_rects:
                img_rect = image_rects[0]
                
                # Analyze position
                page_rect = page.rect
                center_x = (img_rect.x0 + img_rect.x1) / 2
                center_y = (img_rect.y0 + img_rect.y1) / 2
                
                if center_x < page_rect.width * 0.33:
                    h_pos = "left"
                elif center_x > page_rect.width * 0.67:
                    h_pos = "right"
                else:
                    h_pos = "center"
                
                if center_y < page_rect.height * 0.33:
                    v_pos = "top"
                elif center_y > page_rect.height * 0.67:
                    v_pos = "bottom"
                else:
                    v_pos = "middle"
                
                analysis["position"] = f"{v_pos}-{h_pos}"
                
                # Get surrounding text
                expanded_rect = fitz.Rect(
                    max(0, img_rect.x0 - 50),
                    max(0, img_rect.y0 - 50),
                    min(page_rect.width, img_rect.x1 + 50),
                    min(page_rect.height, img_rect.y1 + 50)
                )
                
                surrounding_text = page.get_textbox(expanded_rect).strip()
                if surrounding_text:
                    analysis["surrounding_text"] = surrounding_text[:200]  # Limit length
                
                # Classify image type based on size and position
                img_area = (img_rect.x1 - img_rect.x0) * (img_rect.y1 - img_rect.y0)
                page_area = page_rect.width * page_rect.height
                area_ratio = img_area / page_area
                
                if area_ratio > 0.3:
                    analysis["image_type"] = "large_figure_or_diagram"
                elif area_ratio > 0.1:
                    analysis["image_type"] = "medium_illustration"
                elif h_pos in ["left", "right"] and area_ratio < 0.05:
                    analysis["image_type"] = "inline_icon_or_logo"
                else:
                    analysis["image_type"] = "embedded_figure"
        
        except Exception as e:
            logger.debug(f"Image context analysis failed: {e}")
        
        return analysis
    
    def _extract_tables_with_pdfplumber(self, path: Path) -> List[dict]:
        """Extract tables using pdfplumber for better table detection."""
        table_content = []
        
        try:
            with pdfplumber.open(str(path)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    tables = page.extract_tables()
                    
                    for table_index, table in enumerate(tables):
                        if not table or len(table) < 2:  # Skip empty or single-row tables
                            continue
                        
                        # Convert table to structured text
                        table_text = self._format_table_content(table, page_num, table_index, path.name)
                        
                        table_content.append({
                            "content": table_text,
                            "metadata": {
                                "filename": path.name,
                                "page_number": page_num,
                                "table_index": table_index,
                                "content_type": "pdf_table_structured",
                                "rows": len(table),
                                "columns": len(table[0]) if table else 0
                            }
                        })
        
        except Exception as e:
            logger.debug(f"pdfplumber table extraction failed: {e}")
        
        return table_content
    
    def _format_table_content(self, table: List[List], page_num: int, table_index: int, filename: str) -> str:
        """Format extracted table data into readable text."""
        lines = [f"TABLE {table_index + 1} (Page {page_num}) from {filename}:"]
        
        if not table:
            return lines[0] + "\n[Empty table]"
        
        # Assume first row is header
        headers = table[0] if table else []
        data_rows = table[1:] if len(table) > 1 else []
        
        # Clean headers
        clean_headers = [str(h).strip() if h else f"Column_{i+1}" for i, h in enumerate(headers)]
        lines.append("HEADERS: " + " | ".join(clean_headers))
        
        # Add data rows
        for row_idx, row in enumerate(data_rows[:50]):  # Limit to 50 rows
            clean_row = [str(cell).strip() if cell else "NULL" for cell in row]
            # Pad row to match header length
            while len(clean_row) < len(clean_headers):
                clean_row.append("NULL")
            
            lines.append(f"ROW {row_idx + 1}: " + " | ".join(clean_row[:len(clean_headers)]))
        
        if len(data_rows) > 50:
            lines.append(f"... and {len(data_rows) - 50} more rows")
        
        return "\n".join(lines)
    
    def _extract_tables_with_camelot(self, path: Path) -> List[dict]:
        """Extract complex tables using Camelot (if available)."""
        table_content = []
        
        try:
            import camelot
            
            # Extract tables with different methods
            tables = camelot.read_pdf(str(path), pages='all', flavor='lattice')
            
            if tables.n == 0:
                # Try stream method if lattice fails
                tables = camelot.read_pdf(str(path), pages='all', flavor='stream')
            
            for i, table in enumerate(tables):
                if table.df.empty:
                    continue
                
                # Convert DataFrame to structured text
                table_text = self._format_camelot_table(table, i, path.name)
                
                table_content.append({
                    "content": table_text,
                    "metadata": {
                        "filename": path.name,
                        "page_number": table.page,
                        "table_index": i,
                        "content_type": "pdf_table_camelot",
                        "accuracy": table.accuracy,
                        "rows": len(table.df),
                        "columns": len(table.df.columns)
                    }
                })
        
        except ImportError:
            logger.debug("Camelot not available for table extraction")
        except Exception as e:
            logger.debug(f"Camelot table extraction failed: {e}")
        
        return table_content
    
    def _format_camelot_table(self, table, table_index: int, filename: str) -> str:
        """Format Camelot table to structured text."""
        lines = [f"COMPLEX TABLE {table_index + 1} (Page {table.page}) from {filename}:"]
        lines.append(f"Extraction Accuracy: {table.accuracy:.2f}")
        
        df = table.df
        
        # Add headers
        headers = [f"Col_{i+1}" if col == f"{i}" else str(col) for i, col in enumerate(df.columns)]
        lines.append("HEADERS: " + " | ".join(headers))
        
        # Add data rows
        for idx, row in df.head(50).iterrows():
            row_text = " | ".join([str(val) if pd.notna(val) else "NULL" for val in row.values])
            lines.append(f"ROW {idx + 1}: {row_text}")
        
        if len(df) > 50:
            lines.append(f"... and {len(df) - 50} more rows")
        
        return "\n".join(lines)
    
    def _analyze_document_structure(self, path: Path) -> Optional[dict]:
        """Analyze overall document structure and metadata."""
        try:
            doc = fitz.open(str(path))
            
            analysis = {
                "total_pages": len(doc),
                "metadata": doc.metadata,
                "has_toc": bool(doc.get_toc()),
                "page_sizes": [],
                "text_density": [],
                "image_counts": []
            }
            
            # Analyze each page
            for page_num in range(min(10, len(doc))):  # Sample first 10 pages
                page = doc[page_num]
                
                # Page size
                rect = page.rect
                analysis["page_sizes"].append((rect.width, rect.height))
                
                # Text density
                text = page.get_text()
                text_length = len(text.strip())
                page_area = rect.width * rect.height
                density = text_length / page_area if page_area > 0 else 0
                analysis["text_density"].append(density)
                
                # Image count
                images = page.get_images()
                analysis["image_counts"].append(len(images))
            
            # Calculate averages
            if analysis["text_density"]:
                analysis["avg_text_density"] = np.mean(analysis["text_density"])
                analysis["avg_images_per_page"] = np.mean(analysis["image_counts"])
            
            # Detect document type
            doc_type = self._classify_document_type(analysis)
            
            doc.close()
            
            content = f"""DOCUMENT STRUCTURE ANALYSIS: {path.name}

                Document Metadata:
                • Title: {analysis['metadata'].get('title', 'Not specified')}
                • Author: {analysis['metadata'].get('author', 'Not specified')}
                • Creator: {analysis['metadata'].get('creator', 'Not specified')}
                • Subject: {analysis['metadata'].get('subject', 'Not specified')}

                Document Characteristics:
                • Total Pages: {analysis['total_pages']}
                • Has Table of Contents: {'Yes' if analysis['has_toc'] else 'No'}
                • Average Text Density: {analysis.get('avg_text_density', 0):.2f} chars/area
                • Average Images per Page: {analysis.get('avg_images_per_page', 0):.1f}
                • Detected Document Type: {doc_type}

                Page Size Analysis:
                • Most common page size: {self._get_most_common_page_size(analysis['page_sizes'])}
                • Page size consistency: {'Consistent' if self._is_consistent_page_size(analysis['page_sizes']) else 'Variable'}
                """
            
            return {
                "content": content,
                "metadata": {
                    "filename": path.name,
                    "content_type": "document_structure_analysis",
                    "document_type": doc_type,
                    "total_pages": analysis["total_pages"],
                    "has_toc": analysis["has_toc"]
                }
            }
            
        except Exception as e:
            logger.warning(f"Document structure analysis failed: {e}")
            return None
    
    def _classify_document_type(self, analysis: dict) -> str:
        """Classify document type based on characteristics."""
        avg_text_density = analysis.get("avg_text_density", 0)
        avg_images = analysis.get("avg_images_per_page", 0)
        total_pages = analysis.get("total_pages", 0)
        has_toc = analysis.get("has_toc", False)
        
        if has_toc and total_pages > 20:
            return "book_or_manual"
        elif avg_images > 3 and avg_text_density < 0.01:
            return "presentation_or_visual_document"
        elif avg_images > 1 and avg_text_density > 0.02:
            return "report_with_figures"
        elif avg_text_density > 0.05 and avg_images < 0.5:
            return "text_heavy_document"
        elif total_pages < 5 and avg_text_density > 0.03:
            return "article_or_paper"
        elif avg_images < 0.1 and avg_text_density > 0.04:
            return "academic_paper"
        else:
            return "mixed_content_document"
    
    def _get_most_common_page_size(self, page_sizes: List[Tuple[float, float]]) -> str:
        """Get the most common page size."""
        if not page_sizes:
            return "Unknown"
        
        # Round to avoid floating point issues
        rounded_sizes = [(round(w), round(h)) for w, h in page_sizes]
        
        from collections import Counter
        counter = Counter(rounded_sizes)
        most_common = counter.most_common(1)[0][0]
        
        return f"{most_common[0]}x{most_common[1]}"
    
    def _is_consistent_page_size(self, page_sizes: List[Tuple[float, float]]) -> bool:
        """Check if page sizes are consistent."""
        if len(page_sizes) <= 1:
            return True
        
        first_size = page_sizes[0]
        tolerance = 5  # 5 pixel tolerance
        
        for size in page_sizes[1:]:
            if (abs(size[0] - first_size[0]) > tolerance or 
                abs(size[1] - first_size[1]) > tolerance):
                return False
        
        return True
    
    def _detect_columns(self, text_dict: dict) -> Tuple[bool, int]:
        """Detect if page has column layout."""
        try:
            blocks = text_dict.get("blocks", [])
            if len(blocks) < 2:
                return False, 1
            
            # Get x-coordinates of text blocks
            x_coords = []
            for block in blocks:
                if "lines" in block and block["lines"]:
                    bbox = block.get("bbox", [0, 0, 0, 0])
                    x_coords.append(bbox[0])  # Left edge
            
            if len(x_coords) < 2:
                return False, 1
            
            # Cluster x-coordinates to find columns
            x_coords = sorted(set(x_coords))
            
            # Simple clustering based on gaps
            columns = []
            current_column = [x_coords[0]]
            
            for i in range(1, len(x_coords)):
                if x_coords[i] - x_coords[i-1] > 50:  # Significant gap
                    columns.append(current_column)
                    current_column = [x_coords[i]]
                else:
                    current_column.append(x_coords[i])
            
            columns.append(current_column)
            
            # Consider it columnar if we have 2+ distinct column regions
            if len(columns) >= 2:
                return True, len(columns)
            
            return False, 1
            
        except Exception:
            return False, 1
    
    def _calculate_text_confidence(self, blocks: List[dict]) -> float:
        """Calculate confidence score for text extraction."""
        if not blocks:
            return 0.0
        
        total_chars = 0
        valid_chars = 0
        
        for block in blocks:
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    total_chars += len(text)
                    # Count alphanumeric characters as valid
                    valid_chars += sum(1 for c in text if c.isalnum() or c.isspace())
        
        return valid_chars / total_chars if total_chars > 0 else 0.0
    
    def extract_docx_enhanced(self, path: Path) -> List[dict]:
        """Enhanced extraction for Word documents."""
        content_blocks = []
        
        try:
            doc = Document(str(path))
            
            # Document overview
            overview = self._create_docx_overview(doc, path.name)
            content_blocks.append(overview)
            
            # Extract paragraphs with formatting
            paragraph_content = self._extract_docx_paragraphs(doc, path.name)
            content_blocks.extend(paragraph_content)
            
            # Extract tables
            table_content = self._extract_docx_tables(doc, path.name)
            content_blocks.extend(table_content)
            
            # Extract images
            image_content = self._extract_docx_images(doc, path.name)
            content_blocks.extend(image_content)
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            content_blocks.append(self._create_fallback_content(path, str(e)))
        
        return content_blocks
    
    def _create_docx_overview(self, doc, filename: str) -> dict:
        """Create overview of Word document."""
        try:
            # Count elements
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            # Get core properties
            props = doc.core_properties
            
            content = f"""WORD DOCUMENT OVERVIEW: {filename}

                Document Properties:
                • Title: {props.title or 'Not specified'}
                • Author: {props.author or 'Not specified'}
                • Subject: {props.subject or 'Not specified'}
                • Created: {props.created or 'Not specified'}
                • Modified: {props.modified or 'Not specified'}

                Document Structure:
                • Paragraphs: {paragraph_count}
                • Tables: {table_count}
                • Estimated Pages: {max(1, paragraph_count // 20)}
                """
            
            return {
                "content": content,
                "metadata": {
                    "filename": filename,
                    "content_type": "docx_overview",
                    "paragraph_count": paragraph_count,
                    "table_count": table_count
                }
            }
        except Exception as e:
            logger.warning(f"DOCX overview creation failed: {e}")
            return {
                "content": f"WORD DOCUMENT: {filename} - Overview extraction failed",
                "metadata": {"filename": filename, "content_type": "docx_error"}
            }
    
    def _extract_docx_paragraphs(self, doc, filename: str) -> List[dict]:
        """Extract paragraphs with formatting information."""
        content_blocks = []
        
        current_section = []
        section_count = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Analyze paragraph formatting
            is_heading = self._is_docx_heading(para)
            
            if is_heading and current_section:
                # Save current section
                section_text = "\n".join(current_section)
                if section_text:
                    content_blocks.append({
                        "content": section_text,
                        "metadata": {
                            "filename": filename,
                            "content_type": "docx_section",
                            "section_number": section_count,
                            "paragraph_count": len(current_section)
                        }
                    })
                
                # Start new section
                current_section = [f"HEADING: {text}"]
                section_count += 1
            else:
                current_section.append(text)
        
        # Add final section
        if current_section:
            section_text = "\n".join(current_section)
            content_blocks.append({
                "content": section_text,
                "metadata": {
                    "filename": filename,
                    "content_type": "docx_section",
                    "section_number": section_count,
                    "paragraph_count": len(current_section)
                }
            })
        
        return content_blocks
    
    def _is_docx_heading(self, paragraph) -> bool:
        """Determine if paragraph is a heading."""
        try:
            style_name = paragraph.style.name.lower()
            return ('heading' in style_name or 
                    'title' in style_name or
                    paragraph.style.font.bold)
        except:
            return False
    
    def _extract_docx_tables(self, doc, filename: str) -> List[dict]:
        """Extract tables from Word document."""
        table_content = []
        
        for table_index, table in enumerate(doc.tables):
            try:
                # Convert table to text
                table_text = f"TABLE {table_index + 1} from {filename}:\n"
                
                rows_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace('\n', ' ')
                        row_data.append(cell_text or "NULL")
                    rows_data.append(row_data)
                
                if rows_data:
                    # Assume first row is header
                    headers = rows_data[0]
                    table_text += "HEADERS: " + " | ".join(headers) + "\n"
                    
                    # Add data rows
                    for row_idx, row_data in enumerate(rows_data[1:], 1):
                        table_text += f"ROW {row_idx}: " + " | ".join(row_data) + "\n"
                
                table_content.append({
                    "content": table_text,
                    "metadata": {
                        "filename": filename,
                        "content_type": "docx_table",
                        "table_index": table_index,
                        "rows": len(rows_data),
                        "columns": len(rows_data[0]) if rows_data else 0
                    }
                })
                
            except Exception as e:
                logger.warning(f"Failed to extract table {table_index}: {e}")
        
        return table_content
    
    def _extract_docx_images(self, doc, filename: str) -> List[dict]:
        """Extract images from Word document."""
        image_content = []
        
        try:
            # Extract from document relationships
            rels = doc.part.rels
            image_count = 0
            
            for rel_id, rel in rels.items():
                if "image" in rel.reltype:
                    try:
                        image_data = rel.target_part.blob
                        
                        image_content.append({
                            "content": f"""WORD DOCUMENT IMAGE: {filename}
                            Image {image_count + 1}
                            Format: {rel.target_part.content_type}
                            Size: {len(image_data)} bytes

                            This image was embedded in the Word document and may contain important visual information related to the document content.
                            """,
                            "metadata": {
                                "filename": filename,
                                "content_type": "docx_image",
                                "image_index": image_count,
                                "image_bytes": base64.b64encode(image_data).decode('utf-8'),
                                "content_type_mime": rel.target_part.content_type
                            }
                        })
                        
                        image_count += 1
                        
                    except Exception as e:
                        logger.warning(f"Failed to extract image {image_count}: {e}")
        
        except Exception as e:
            logger.warning(f"DOCX image extraction failed: {e}")
        
        return image_content
    
    def extract_epub_enhanced(self, path: Path) -> List[dict]:
        """Enhanced extraction for EPUB files."""
        content_blocks = []
        
        try:
            with zipfile.ZipFile(str(path), 'r') as epub_zip:
                # Parse EPUB structure
                content_opf = self._find_epub_content_opf(epub_zip)
                if not content_opf:
                    return [self._create_fallback_content(path, "Could not find EPUB content.opf")]
                
                # Extract metadata
                metadata = self._extract_epub_metadata(epub_zip, content_opf)
                content_blocks.append(metadata)
                
                # Extract chapters
                chapters = self._extract_epub_chapters(epub_zip, content_opf, path.name)
                content_blocks.extend(chapters)
                
        except Exception as e:
            logger.error(f"EPUB extraction failed: {e}")
            content_blocks.append(self._create_fallback_content(path, str(e)))
        
        return content_blocks
    
    def _find_epub_content_opf(self, epub_zip) -> Optional[str]:
        """Find the content.opf file in EPUB."""
        try:
            # Check META-INF/container.xml
            container_xml = epub_zip.read('META-INF/container.xml')
            root = ET.fromstring(container_xml)
            
            # Find rootfile
            for rootfile in root.findall('.//{urn:oasis:names:tc:opendocument:xmlns:container}rootfile'):
                return rootfile.get('full-path')
        except:
            pass
        
        # Fallback: look for common locations
        common_paths = ['OEBPS/content.opf', 'content.opf', 'OPS/content.opf']
        for path in common_paths:
            if path in epub_zip.namelist():
                return path
        
        return None
    
    def _extract_epub_metadata(self, epub_zip, content_opf_path: str) -> dict:
        """Extract EPUB metadata."""
        try:
            content_opf = epub_zip.read(content_opf_path).decode('utf-8')
            root = ET.fromstring(content_opf)
            
            # Extract metadata
            metadata = {}
            ns = {'dc': 'http://purl.org/dc/elements/1.1/'}
            
            title = root.find('.//dc:title', ns)
            creator = root.find('.//dc:creator', ns)
            description = root.find('.//dc:description', ns)
            language = root.find('.//dc:language', ns)
            
            content = f"""EPUB METADATA:
                Title: {title.text if title is not None else 'Unknown'}
                Author: {creator.text if creator is not None else 'Unknown'}
                Description: {description.text if description is not None else 'No description'}
                Language: {language.text if language is not None else 'Unknown'}
                """
            
            return {
                "content": content,
                "metadata": {
                    "content_type": "epub_metadata",
                    "title": title.text if title is not None else 'Unknown',
                    "author": creator.text if creator is not None else 'Unknown'
                }
            }
            
        except Exception as e:
            logger.warning(f"EPUB metadata extraction failed: {e}")
            return {
                "content": "EPUB METADATA: Extraction failed",
                "metadata": {"content_type": "epub_metadata_error"}
            }
    
    def _extract_epub_chapters(self, epub_zip, content_opf_path: str, filename: str) -> List[dict]:
        """Extract chapters from EPUB."""
        chapters = []
        
        try:
            content_opf = epub_zip.read(content_opf_path).decode('utf-8')
            root = ET.fromstring(content_opf)
            
            # Get base directory
            base_dir = '/'.join(content_opf_path.split('/')[:-1])
            if base_dir:
                base_dir += '/'
            
            # Find spine items
            spine_items = []
            for itemref in root.findall('.//{http://www.idpf.org/2007/opf}itemref'):
                idref = itemref.get('idref')
                if idref:
                    spine_items.append(idref)
            
            # Get manifest items
            manifest_items = {}
            for item in root.findall('.//{http://www.idpf.org/2007/opf}item'):
                item_id = item.get('id')
                href = item.get('href')
                if item_id and href:
                    manifest_items[item_id] = base_dir + href
            
            # Extract chapter content
            for chapter_num, spine_id in enumerate(spine_items[:20], 1):  # Limit to 20 chapters
                if spine_id in manifest_items:
                    chapter_path = manifest_items[spine_id]
                    
                    try:
                        chapter_content = epub_zip.read(chapter_path).decode('utf-8')
                        
                        # Parse HTML and extract text
                        soup = BeautifulSoup(chapter_content, 'html.parser')
                        
                        # Remove script and style elements
                        for script in soup(["script", "style"]):
                            script.decompose()
                        
                        text = soup.get_text()
                        # Clean up text
                        lines = (line.strip() for line in text.splitlines())
                        text = '\n'.join(line for line in lines if line)
                        
                        if text and len(text.strip()) > 100:  # Only include substantial chapters
                            chapters.append({
                                "content": f"CHAPTER {chapter_num} from {filename}:\n\n{text}",
                                "metadata": {
                                    "filename": filename,
                                    "content_type": "epub_chapter",
                                    "chapter_number": chapter_num,
                                    "chapter_path": chapter_path,
                                    "word_count": len(text.split())
                                }
                            })
                    
                    except Exception as e:
                        logger.warning(f"Failed to extract chapter {chapter_num}: {e}")
        
        except Exception as e:
            logger.warning(f"EPUB chapter extraction failed: {e}")
        
        return chapters
    
    def _create_document_overview(self, doc, filename: str) -> dict:
        """Create document overview."""
        try:
            metadata = doc.metadata
            page_count = len(doc)
            
            content = f"""PDF DOCUMENT OVERVIEW: {filename}

            Document Metadata:
            • Title: {metadata.get('title', 'Not specified')}
            • Author: {metadata.get('author', 'Not specified')}
            • Creator: {metadata.get('creator', 'Not specified')}
            • Producer: {metadata.get('producer', 'Not specified')}
            • Subject: {metadata.get('subject', 'Not specified')}

            Document Statistics:
            • Total Pages: {page_count}
            • File Size: {Path(doc.name).stat().st_size / (1024*1024):.2f} MB
            • Creation Date: {metadata.get('creationDate', 'Not specified')}
            • Modification Date: {metadata.get('modDate', 'Not specified')}
            """
            
            return {
                "content": content,
                "metadata": {
                    "filename": filename,
                    "content_type": "pdf_document_overview",
                    "page_count": page_count,
                    "has_metadata": bool(any(metadata.values()))
                }
            }
        except Exception as e:
            logger.warning(f"Document overview creation failed: {e}")
            return {
                "content": f"PDF DOCUMENT: {filename} - Overview extraction failed",
                "metadata": {"filename": filename, "content_type": "pdf_overview_error"}
            }
    
    def _create_fallback_pdf_content(self, path: Path, error: str) -> List[dict]:
        """Create fallback content for PDF processing failures."""
        return [{
            "content": f"PDF DOCUMENT: {path.name} - Processing failed: {error}",
            "metadata": {
                "filename": path.name,
                "content_type": "pdf_processing_error",
                "error": error
            }
        }]
    
    def _create_fallback_content(self, path: Path, error: str) -> dict:
        """Create fallback content for general processing failures."""
        return {
            "content": f"DOCUMENT: {path.name} - Processing failed: {error}",
            "metadata": {
                "filename": path.name,
                "content_type": "processing_error",
                "error": error
            }
        }

# Global processor instance
_document_processor = None

def get_document_processor():
    """Get or create document processor instance."""
    global _document_processor
    if _document_processor is None:
        _document_processor = EnhancedDocumentProcessor()
    return _document_processor

def extract_pdf_enhanced(path: Path) -> List[dict]:
    """Enhanced PDF extraction function."""
    processor = get_document_processor()
    return processor.extract_pdf_enhanced(path)

def extract_docx_enhanced(path: Path) -> List[dict]:
    """Enhanced DOCX extraction function."""
    processor = get_document_processor()
    return processor.extract_docx_enhanced(path)

def extract_epub_enhanced(path: Path) -> List[dict]:
    """Enhanced EPUB extraction function."""
    processor = get_document_processor()
    return processor.extract_epub_enhanced(path)
